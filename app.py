import streamlit as st
import feedparser
from datetime import datetime, timedelta
import openai
from urllib.parse import quote
import time
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from io import BytesIO
import docx

# 1. 페이지 설정
st.set_page_config(page_title="Samsung Securities FO Strategy", page_icon="💙", layout="wide")

# 사이드바 설정 (Secrets 로드)
with st.sidebar:
    st.title("⚙️ Strategy Settings")
    if "OPENAI_API_KEY" in st.secrets:
        api_key = st.secrets["OPENAI_API_KEY"]
        st.success("✅ API Key 로드 완료")
    else:
        api_key = st.text_input("OpenAI API Key를 입력하세요", type="password")
    
    news_count = st.slider("최대 분석 뉴스 개수", 1, 10, 3)

# --- 워드 파일 하이퍼링크 삽입을 위한 헬퍼 함수 ---
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    c = docx.oxml.shared.OxmlElement('w:rStyle')
    c.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')
    rPr.append(c)
    new_run.append(rPr)
    text_obj = docx.oxml.shared.OxmlElement('w:t')
    text_obj.text = text
    new_run.append(text_obj)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# --- 워드 리포트 생성 함수 ---
def create_word_report(report_items):
    doc = Document()
    doc.add_heading('Samsung Securities FO Strategic Intelligence Report', 0)
    doc.add_paragraph(f"Generated Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("-" * 80)

    for item in report_items:
        # 워드 대제목: 제목 (언론사, 발간날짜)
        doc.add_heading(f"{item['title']} ({item['source']}, {item['date']})", level=1)
        
        p = doc.add_paragraph("Original Article: ")
        add_hyperlink(p, item['link'], "Click here to read the full article")
        
        doc.add_heading('Strategic Analysis', level=2)
        doc.add_paragraph(item['analysis'])
        doc.add_page_break()
    
    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# 2. 뉴스 크롤링 함수
def get_news(keyword):
    sources_list = ["JPMorgan", "Goldman Sachs", "UBS", "Morgan Stanley", "Bloomberg", "Reuters", "Financial Times", "Forbes", "WSJ", "Blackstone", "KKR", "BlackRock", "Julius Baer", "Pictet", "Lombard Odier"]
    full_query = f'("{keyword}") AND ({" OR ".join([f'"{s}"' for s in sources_list])})'
    
    def fetch(query):
        encoded = quote(query)
        rss_url = f"https://news.google.com/rss/search?q={encoded}+when:7d&hl=en-US&gl=US&ceid=US:en"
        feed = feedparser.parse(rss_url)
        now = datetime.now()
        filtered = []
        for entry in feed.entries:
            try:
                p_date = datetime(*(entry.published_parsed[:6]))
                if now - p_date <= timedelta(days=7):
                    filtered.append(entry)
            except: filtered.append(entry)
        return filtered

    res = fetch(full_query)
    if not res: res = fetch(f'"{keyword}"')
    return res[:news_count]

# 3. AI 분석 함수 (요약 분량 대폭 강화)
def analyze_article(title, link):
    if not api_key: return "API Key Error"
    client = openai.OpenAI(api_key=api_key)
    prompt = f"""
    당신은 삼성증권 전략기획실 및 SNI(Samsung Private Banking) 본부의 수석 전략 컨설턴트입니다. 
    다음 뉴스를 바탕으로 심층 전략 분석 보고서를 작성하세요.

    뉴스 제목: {title}

    [보고서 양식 및 지침]
    1. 요약: 
       - 이 뉴스의 핵심 내용, 발생 배경, 주요 데이터 및 글로벌 시장의 거시적 맥락을 포함하여 **최소 20문장에서 30문장 사이**로 아주 상세하게 기술하세요. 
       - 기사의 모든 중요한 세부 사항을 놓치지 말고 리포트 형식으로 정리하세요.

    2. 전략 아이디어: 
       ① [전략]: 삼성증권 SNI가 즉각적 혹은 장기적으로 실행해야 할 구체적인 비즈니스 모델이나 서비스 혁신 방안.
       ② [선정이유 및 전략적 타당성]: 글로벌 리딩 뱅크의 사례와 국내 UHNWI 시장(상속, 기업 승계, 대체투자 수요 등)의 특수성을 결합한 깊이 있는 분석.
       ③ [기대 효과 및 향후 과제]: 예상되는 정량적/정성적 성과와 실행을 위한 구체적인 로드맵 및 선결 과제.
    
    언어: 한국어 (매우 전문적이고 격식 있는 비즈니스 문체 사용)
    """
    response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}], temperature=0.5)
    return response.choices[0].message.content

# 4. 메인 UI 및 로직
st.title("🏛️ Samsung Securities FO Strategic Intel")
st.write("글로벌 최신 뉴스 기반 삼성증권 SNI 심층 전략 리포트")

if st.button('🚀 분석 및 리포트 생성'):
    if not api_key:
        st.error("OpenAI API Key를 설정해 주세요.")
    else:
        with st.spinner('글로벌 금융 데이터 분석 및 상세 리포트 작성 중...'):
            news_items = get_news("family office")
            if not news_items:
                st.warning("최근 7일 이내 뉴스가 없습니다.")
            else:
                all_reports = []
                for item in news_items:
                    # 데이터 가공
                    source_name = getattr(item, 'source', {}).get('title', 'Global Source')
                    try:
                        dt = datetime(*(item.published_parsed[:6]))
                        formatted_date = dt.strftime('%Y-%m-%d')
                    except:
                        formatted_date = "N/A"
                        
                    analysis_result = analyze_article(item.title, item.link)
                    
                    # 화면 표시: ## 제목 (언론사, 발간날짜)
                    with st.container():
                        st.markdown(f"## {item.title} ({source_name}, {formatted_date})")
                        st.write(f"🔗 [기사 원문 링크]({item.link})")
                        st.markdown(analysis_result)
                        st.divider()
                    
                    all_reports.append({
                        'title': item.title,
                        'source': source_name,
                        'link': item.link,
                        'date': formatted_date,
                        'analysis': analysis_result
                    })
                    time.sleep(1) # API 안정성 확보

                # 워드 파일 생성
                docx_file = create_word_report(all_reports)
                
                st.download_button(
                    label="📂 워드 파일(.docx) 다운로드",
                    data=docx_file,
                    file_name=f"Samsung_FO_Strategy_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
